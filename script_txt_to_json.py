import os
import json
from docx import Document
from openai import OpenAI

# 📌 ודא שיש לך משתנה סביבה OPENAI_API_KEY
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# ⚙️ פרטי המשתמש – דוגמה ברורה
user_input = {
    "business_name": "מאפיית חלום",
    "business_type": "bakery",   # אפשרויות: cafe, food_truck, restaurant, bar, bakery, catering
    "area_sqm": 80,
    "seating_capacity": 12,
    "employees": 5,
    "city": "תל אביב",
    "has_gas": True,
    "serves_meat": False,
    "has_delivery": True,
    "has_alcohol": False
}

# 📝 סוגי העסקים האפשריים
BUSINESS_TYPES = ["cafe", "food_truck", "restaurant", "bar", "bakery", "catering"]

# 📝 קטגוריות אפשריות
CATEGORIES = [
    "בריאות ותברואה",
    "בטיחות אש",
    "רישוי ותכנון",
    "סביבה ופסולת",
    "תפעול וניהול עובדים",
    "אחר"
]

# 📝 טבלה של שדות אפשריים לחוקים
FIELDS_TABLE = """
שדות אפשריים לסיווג חוק:
- business_type: cafe, food_truck, restaurant, bar, bakery, catering
- has_gas: true / false
- serves_meat: true / false
- has_delivery: true / false
- has_alcohol: true / false
- area_sqm: מספר (טווחים לדוגמה: מתחת ל-50, מעל 200)
- seating_capacity: מספר (טווחים לדוגמה: עד 20, מעל 100)
- employees: מספר (טווחים לדוגמה: מעל 10 עובדים)
- city: אפשר להתנות בעיר מסוימת, אם החוק קשור לרשות מקומית
"""

def extract_text_from_docx(docx_path):
    """קריאת כל הטקסט והטבלאות מקובץ Word"""
    doc = Document(docx_path)
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                full_text.append(" | ".join(row_data))

    return "\n".join(full_text)

def convert_rules_with_ai(text, start_id=1):
    """שולח חלק טקסט ל-ChatGPT ומחזיר JSON"""
    prompt = f"""
אתה מקבל קטע מתוך קובץ עם חוקים ודרישות.

נתוני העסק לדוגמה:
{json.dumps(user_input, ensure_ascii=False, indent=2)}

סוגי עסקים אפשריים:
{BUSINESS_TYPES}

קטגוריות אפשריות:
{CATEGORIES}

{FIELDS_TABLE}

המבנה של כל חוק:
{{
  "id": "RXXX",
  "title": "שם החוק",
  "category": "בריאות ותברואה / בטיחות אש / רישוי ותכנון / סביבה ופסולת / תפעול וניהול עובדים / אחר",
  "applies_when": {{
    "business_type": ["restaurant", "bar"],
    "has_gas": [true, false],
    "serves_meat": [true, false],
    "has_delivery": [true, false],
    "has_alcohol": [true, false],
    "min_area": null,
    "max_area": null,
    "seating_capacity": null,
    "employees": null,
    "city": null
  }},
  "actions": ["פעולה 1", "פעולה 2"],
  "priority": "קריטי/גבוה/בינוני/נמוך",
  "estimated_cost": "עלות משוערת (₪ או 'ללא עלות נוספת')"
}}

דגשים חשובים:
- סווג כל חוק לקטגוריה רלוונטית אחת מתוך הרשימה.
- ייתכן שחוק יתאים ליותר מסוג עסק אחד → רשום את כולם.
- אם החוק תלוי בגז/בשר/אלכוהול/משלוחים → ציין זאת ב-applies_when.
- אם זה כלל כללי (כמו תליית רישיון) → business_type = כל {BUSINESS_TYPES}.
- הוסף עלות משוערת ריאלית (או "ללא עלות נוספת").
- חובה למספר ברצף (R0001, R0002...).
- אל תסתפק בדוגמה אחת, תסווג לפי ההיגיון שלך.
- החזר אך ורק JSON תקין.
- אם אין חוקים בקטע → החזר {{"rules": []}} בלבד.

קטע מתוך הקובץ:
{text}
    """

    response = client.chat.completions.create(
        model="gpt-4.1",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )

    ai_text = response.choices[0].message.content
    return json.loads(ai_text)

def split_text(text, chunk_size=5000):
    """פיצול הטקסט לחתיכות קטנות יותר"""
    return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]

if __name__ == "__main__":
    docx_file = "18-07-2022_4.2A.docx"
    output_file = "rules.json"

    print("📂 קורא את הקובץ...")
    text = extract_text_from_docx(docx_file)

    chunks = split_text(text, chunk_size=5000)
    print(f"✂️ הקובץ פוצל ל-{len(chunks)} חלקים")

    # ✂️ עיבוד רק חצי קובץ
    half_index = max(1, len(chunks) // 2)
    chunks = chunks[:half_index]
    print(f"📂 מעבד רק {len(chunks)} חלקים (חצי קובץ)")

    all_rules = {"rules": []}
    current_id = 1

    for i, chunk in enumerate(chunks, start=1):
        print(f"🤖 שולח ל-ChatGPT חלק {i}/{len(chunks)}...")
        try:
            ai_data = convert_rules_with_ai(chunk, start_id=current_id)

            rules = ai_data.get("rules", [])
            for j, rule in enumerate(rules, start=0):
                rule["id"] = f"R{current_id + j:04d}"
            current_id += len(rules)

            all_rules["rules"].extend(rules)
        except Exception as e:
            print(f"❌ שגיאה בחלק {i}: {e}")

    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(all_rules, f, ensure_ascii=False, indent=2)

    print(f"✅ קובץ JSON נוצר בהצלחה: {output_file}")
